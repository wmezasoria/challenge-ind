import PyPDF2

from docx import Document
from docx.shared import Inches
from docx.shared import Pt

from docx.shared import RGBColor
from docx.oxml.ns import qn

from openai import OpenAI
from docx import Document
from dotenv import dotenv_values


# from reportlab.lib.pagesizes import letter
# from reportlab.platypus import SimpleDocTemplate, Paragraph, PageTemplate, Frame
# from reportlab.lib.styles import getSampleStyleSheet

config = dotenv_values(".env")
client = OpenAI(api_key=config["OPENAI_API_KEY"])


cv = ""

# Abrir el archivo PDF en modo de lectura binaria
with open('archivo.pdf', 'rb') as archivo_pdf:
    # Crear un objeto PdfReader
    lector_pdf = PyPDF2.PdfReader(archivo_pdf)
    
    # Obtener el número de páginas en el PDF
    num_paginas = len(lector_pdf.pages)
    
    # Iterar a través de cada página del PDF
    for pagina_numero in range(num_paginas):
        # Obtener el objeto Page correspondiente a la página actual
        pagina = lector_pdf.pages[pagina_numero]
        
        # Extraer el texto de la página
        cv += pagina.extract_text()
        
        # Imprimir el texto extraído
        #cv = texto
#print(texto)
    
    #print(cv)
res = client.chat.completions.create(
    #model="gpt-4",  # Selecciona el modelo GPT-3
    model="gpt-3.5-turbo",
    messages=[
        {"role": "system", "content": "eres el asistente de recursos humanos."},
        {"role": "user", "content": "dime los nombres y apellidos del postulante en mayusculas, ejemplo MEZA SORIA WILLIAM JOSE: "},
        {"role": "user", "content": cv}

    ]

)
print(res.choices[0].message.content)
nombres_postulante = res.choices[0].message.content
#print(texto)

res_resumen = client.chat.completions.create(
    model="gpt-3.5-turbo",
    messages=[
        {"role": "system", "content": "eres el asistente de recursos humanos."},
        #{"role": "user", "content": "genere resúmenes personalizados separado por las siguietntes secciones: <RESUMEN>, <FORMACION> <CURSOS> <TECNOLOGIAS> <EXPERIENCIA> y extraiga insights estratégicos valiosos del siguiente texto: "},
        {"role": "user", "content": "dime un resumen de la siguiente hoja de vida, en no mas de 50 palabras "},
        {"role": "user", "content": cv}

    ]

)



res_nacionalidad = client.chat.completions.create(
    model="gpt-3.5-turbo",
    messages=[
        {"role": "system", "content": "eres el asistente de recursos humanos."},
        {"role": "user", "content": "dime en maximo 2 palabras la nacionalidad del postulante, ejemplo: PERUANA, ECUATORIANA "},
        {"role": "user", "content": cv}

    ]

)


nacionalidad = res_nacionalidad.choices[0].message.content

resumen = res_resumen.choices[0].message.content



# Crear un nuevo documento Word
document = Document()

# Agregar un encabezado 
#document.add_heading('William', 0)
heading = document.add_heading('', 0)
heading1 = document.add_heading('', 1)
# heading = document.add_heading('\n '+nacionalidad, 0)

heading.add_run(nombres_postulante).bold = True
heading1.add_run('\n'+nacionalidad)

heading.style.font.name = 'Calibri'
heading.style.font.size = Pt(18)
heading1.style.font.name = 'Calibri'
heading1.style.font.size = Pt(11)


#heading_paragraph = heading.paragraph_format + heading1.paragraph_format
# shading_elm = heading_paragraph._element.get_or_add_tcPr().get_or_add_shd()
# shading_elm.set(qn('w:fill'), 'blue')



subtitle = document.add_heading(level=2)
subtitle_run = subtitle.add_run('RESUMEN')
document.add_paragraph(resumen)

# Crear una tabla de 1 fila y 2 columnas
table = document.add_table(rows=1, cols=2)


subtitle = document.add_heading(level=2)
subtitle_run = subtitle.add_run('FORMACION')
# Agregar contenido a la columna izquierda
table.cell(0, 0).text = 'Columna Izquierda'
table.cell(0, 0).add_paragraph('Este es el contenido de la columna izquierda.')
table.cell(0, 0).add_paragraph('Más texto en la columna izquierda.')

# Agregar contenido a la columna derecha
table.cell(0, 1).text = 'Columna Derecha'
table.cell(0, 1).add_paragraph('Aquí va el contenido de la columna derecha.')
table.cell(0, 1).add_paragraph('Test')

# Ajustar el ancho de las columnas
table.columns[0].width = Inches(3.25)
table.columns[1].width = Inches(3.25)


subtitle1 = document.add_heading(level=2)
subtitle_run = subtitle1.add_run('CURSOS Y CERTIFICACIONES')
subtitle2 = document.add_heading(level=2)
subtitle_run = subtitle2.add_run('TECNOLOGIAS')
subtitle3 = document.add_heading(level=2)
subtitle_run = subtitle3.add_run('EXPERIENCIA')


# Guardar el documento Word
document.save('main.docx')

